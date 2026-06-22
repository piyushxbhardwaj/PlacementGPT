import os
from typing import List, Dict, Any
from pypdf import PdfReader
import docx

class DocumentParser:
    @staticmethod
    def parse(file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Parses a document and returns a list of dictionaries with page numbers and text content.
        List[dict] -> [{"page_number": int, "text": str}, ...]
        """
        ext = file_type.lower() or os.path.splitext(file_path)[1].lower().replace(".", "")
        
        if ext == "pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext in ["docx", "doc"]:
            return DocumentParser._parse_docx(file_path)
        elif ext in ["txt", "md"]:
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append({
                    "page_number": i + 1,
                    "text": text.strip()
                })
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF file {file_path}: {str(e)}")
        return pages

    @staticmethod
    def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
        try:
            doc = docx.Document(file_path)
            # docx text doesn't have reliable page numbers, extract all text as page 1
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables as well for completeness
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return [{
                "page_number": 1,
                "text": "\n".join(full_text)
            }]
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX file {file_path}: {str(e)}")

    @staticmethod
    def _parse_txt(file_path: str) -> List[Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return [{
                "page_number": 1,
                "text": content.strip()
            }]
        except Exception as e:
            raise RuntimeError(f"Error parsing TXT file {file_path}: {str(e)}")
